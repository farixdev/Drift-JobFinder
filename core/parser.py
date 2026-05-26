import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

SUPPORTED_EXTENSIONS = (".pdf", ".doc", ".docx")


def extract(file_path: str) -> str:
    path = Path(file_path)
    if not path.is_file():
        raise ValueError(f"File not found: {file_path}")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Use PDF, DOC, or DOCX."
        )

    if ext == ".pdf":
        text = _extract_pdf(path)
    elif ext == ".docx":
        text = _extract_docx(path)
    else:
        text = _extract_doc(path)

    text = _normalize_text(text)
    if len(text) < 20:
        raise ValueError(
            "Could not extract enough readable text from this resume. "
            "If it is a scanned PDF, export it as a text-based PDF or DOCX."
        )
    return text


def _normalize_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_pdf(path: Path) -> str:
    errors: list[str] = []

    try:
        text = _pdf_pdfplumber(path)
        if text:
            return text
    except Exception as exc:
        errors.append(f"pdfplumber: {exc}")

    try:
        text = _pdf_pymupdf(path)
        if text:
            return text
    except Exception as exc:
        errors.append(f"pymupdf: {exc}")

    try:
        text = _pdf_pypdf(path)
        if text:
            return text
    except Exception as exc:
        errors.append(f"pypdf: {exc}")

    detail = "; ".join(errors[:2]) if errors else "no text in file"
    raise ValueError(f"Could not read PDF ({detail}).")


def _pdf_pdfplumber(path: Path) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            chunk = page.extract_text() or ""
            if not chunk.strip():
                tables = page.extract_tables() or []
                for table in tables:
                    for row in table:
                        cells = [c for c in row if c]
                        if cells:
                            parts.append(" ".join(str(c) for c in cells))
            else:
                parts.append(chunk)
    return _normalize_text("\n".join(parts))


def _pdf_pymupdf(path: Path) -> str:
    import fitz

    parts: list[str] = []
    with fitz.open(str(path)) as doc:
        for page in doc:
            parts.append(page.get_text("text") or "")
    return _normalize_text("\n".join(parts))


def _pdf_pypdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return _normalize_text("\n".join(parts))


def _extract_docx(path: Path) -> str:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError

    try:
        doc = Document(str(path))
    except PackageNotFoundError:
        if _is_zip_docx(path):
            raise ValueError("File looks corrupted or is not a valid DOCX.")
        raise

    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return _normalize_text("\n".join(parts))


def _extract_doc(path: Path) -> str:
    if _is_zip_docx(path):
        return _extract_docx(path)

    for extractor in (_extract_doc_win32, _extract_doc_libreoffice, _extract_doc_antiword):
        try:
            text = extractor(path)
            if _normalize_text(text):
                return text
        except Exception:
            continue

    raise ValueError(
        "Could not read .doc file. Save as .docx or PDF, or install Microsoft Word "
        "(Windows) / LibreOffice."
    )


def _is_zip_docx(path: Path) -> bool:
    try:
        with path.open("rb") as handle:
            return handle.read(2) == b"PK"
    except OSError:
        return False


def _extract_doc_win32(path: Path) -> str:
    if sys.platform != "win32":
        return ""
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = None
    try:
        doc = word.Documents.Open(str(path.resolve()), ReadOnly=True)
        text = doc.Content.Text or ""
        return text.replace("\r", "\n")
    finally:
        if doc is not None:
            doc.Close(False)
        word.Quit()


def _extract_doc_libreoffice(path: Path) -> str:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice and sys.platform == "win32":
        for candidate in (
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ):
            if Path(candidate).is_file():
                soffice = candidate
                break
    if not soffice:
        return ""

    with tempfile.TemporaryDirectory() as tmp:
        out_dir = Path(tmp)
        result = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "txt:Text",
                "--outdir",
                str(out_dir),
                str(path.resolve()),
            ],
            capture_output=True,
            text=True,
            timeout=60,
            check=False,
        )
        if result.returncode != 0:
            return ""
        txt_files = list(out_dir.glob("*.txt"))
        if not txt_files:
            return ""
        return txt_files[0].read_text(encoding="utf-8", errors="replace")


def _extract_doc_antiword(path: Path) -> str:
    antiword = shutil.which("antiword")
    if not antiword:
        return ""
    result = subprocess.run(
        [antiword, str(path)],
        capture_output=True,
        text=True,
        timeout=30,
        check=False,
    )
    if result.returncode != 0:
        return ""
    return result.stdout
